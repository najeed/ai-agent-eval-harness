import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def create_security_roadmap():
    doc = Document()

    # --- Document Styling ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # --- Title Page ---
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Security & Compliance Roadmap")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(27, 59, 95)  # #1B3B5F - AgentV Primary Blue

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("The Path to SOC 2 Type I Certification (Q4 2026)")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(5):
        doc.add_paragraph()

    # Version & Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version: 1.0.0\nDate: May 6, 2026\nOwner: Security & Governance Team")
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("1. Executive Summary")
    doc.add_paragraph("2. Current Security Posture")
    doc.add_paragraph("3. The Road to SOC 2 Type I (Q4 2026)")
    doc.add_paragraph("4. Quarterly Milestones")
    doc.add_paragraph("   4.1 Q2 2026: Foundation & Hardening")
    doc.add_paragraph("   4.2 Q3 2026: Readiness & Audit Prep")
    doc.add_paragraph("   4.3 Q4 2026: Certification & Public Trust")
    doc.add_paragraph("5. Framework Alignment (NIST, ISO, HIPAA)")
    doc.add_paragraph("6. Resource Requirements")
    doc.add_paragraph("7. Conclusion")

    doc.add_page_break()

    # --- 1. Executive Summary ---
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "AgentV is committed to maintaining the highest standards of security and data integrity. "
        "As autonomous agents handle increasingly sensitive enterprise data, proving the security "
        "of the verification platform itself is paramount. This roadmap outlines our strategic "
        "journey toward achieving SOC 2 Type I certification by Q4 2026, ensuring our customers "
        "can trust AgentV with their most critical business logic."
    )

    # --- 2. Current Security Posture ---
    doc.add_heading("2. Current Security Posture", level=1)
    doc.add_paragraph("AgentV v1.6.0 already incorporates several 'Secure-by-Design' principles:")
    doc.add_paragraph(
        "Identity Registry: Centralized management of Ed25519 cryptographic keys.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "PBAC Governance: Permissions-Based Access Control for tool and resource isolation.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "PII Redaction: Automated scanning and redaction of 16+ sensitive data patterns.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Forensic DNA: Write-Once-Read-Many (WORM) audit logs and SHA-256 artifact hashing.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "NIST Alignment: Explicit alignment with NIST AI-100-1 (AI RMF).", style="List Bullet"
    )

    # --- 3. The Road to SOC 2 Type I (Q4 2026) ---
    doc.add_heading("3. The Road to SOC 2 Type I (Q4 2026)", level=1)
    doc.add_paragraph(
        "Our journey to SOC 2 Type I will focus on the following Trust Services Criteria (TSC):"
    )
    doc.add_paragraph(
        "Security: Protection against unauthorized access, use, or modification.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Confidentiality: Protection of data designated as confidential.", style="List Bullet"
    )
    doc.add_paragraph(
        "Availability: Ensuring the system is available for operation and use as committed.",
        style="List Bullet",
    )

    # --- 4. Quarterly Milestones ---
    doc.add_heading("4. Quarterly Milestones", level=1)

    doc.add_heading("4.1 Q2 2026: Foundation & Hardening", level=2)
    doc.add_paragraph(
        "Focus: Strengthening the technical foundation and formalizing core security policies."
    )
    doc.add_paragraph(
        "Implement automated vulnerability scanning in CI/CD pipelines (SAST/DAST).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Formalize the Incident Response Plan (IRP) and conduct first tabletop exercise.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Enhance PBAC 2.0 with namespace-specific write permissions and audit logging.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Establish the Data Classification Policy and Inventory.", style="List Bullet"
    )

    doc.add_heading("4.2 Q3 2026: Readiness & Audit Prep", level=2)
    doc.add_paragraph(
        "Focus: Bridging the gap between technical features and formal audit requirements."
    )
    doc.add_paragraph(
        "Conduct a formal SOC 2 Gap Analysis with a third-party consultant.", style="List Bullet"
    )
    doc.add_paragraph(
        "Automate evidence collection by linking VC v3 certificates to the GRC platform.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Roll out mandatory security awareness training for all contributors.", style="List Bullet"
    )
    doc.add_paragraph(
        "Perform an internal security audit of the Identity Registry and Vaulting system.",
        style="List Bullet",
    )

    doc.add_heading("4.3 Q4 2026: Certification & Public Trust", level=2)
    doc.add_paragraph("Focus: Final audit execution and public certification.")
    doc.add_paragraph(
        "Execute the SOC 2 Type I audit with a certified CPA firm.", style="List Bullet"
    )
    doc.add_paragraph(
        "Remediate any findings identified during the observation period.", style="List Bullet"
    )
    doc.add_paragraph(
        "Launch the AgentV Trust Portal for customer-facing security documentation.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Achieve SOC 2 Type I Certification by December 31, 2026.", style="List Bullet"
    )

    # --- 5. Framework Alignment ---
    doc.add_heading("5. Framework Alignment", level=1)
    doc.add_paragraph(
        "While SOC 2 is the primary objective, our roadmap ensures simultaneous alignment with:"
    )
    doc.add_paragraph(
        "NIST AI-100-1: Continued leadership in AI Risk Management.", style="List Bullet"
    )
    doc.add_paragraph(
        "ISO 27001: Leveraging SOC 2 controls for future ISO certification.", style="List Bullet"
    )
    doc.add_paragraph(
        "HIPAA/GDPR: Releasing dedicated compliance packs for regulated industries.",
        style="List Bullet",
    )

    # --- 6. Resource Requirements ---
    doc.add_heading("6. Resource Requirements", level=1)
    doc.add_paragraph(
        "To achieve these milestones, we will allocate resources across Security Engineering, "
        "Legal/Compliance, and DevOps, including the procurement of a dedicated GRC "
        "automation platform."
    )

    # --- 7. Conclusion ---
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "The journey to SOC 2 Type I is a critical milestone in AgentV's mission to become the "
        "authoritative verification layer for enterprise AI. By Q4 2026, AgentV will provide not "
        "just technical verification, but a certified trust framework that meets the demands "
        "of the most regulated industries."
    )

    # --- Footer ---
    section = doc.sections[0]
    footer = section.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("© 2026 AgentV. All Rights Reserved. Confidential & Proprietary.")
    run.font.size = Pt(8)

    # --- Save Document ---
    output_path = "Security_and_Compliance_Roadmap.docx"
    doc.save(output_path)
    print(f"Roadmap generated at: {os.path.abspath(output_path)}")


if __name__ == "__main__":
    create_security_roadmap()
